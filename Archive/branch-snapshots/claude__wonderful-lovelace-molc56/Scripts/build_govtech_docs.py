#!/usr/bin/env python3
"""
build_govtech_docs.py

Builds all strategy paper outputs from Markdown source files.

Output formats:
  - HTML:  generated from Markdown using Python-Markdown + accessible CSS
  - DOCX:  generated using python-docx (or mammoth if available)
  - PDF:   generated via weasyprint from HTML (optional)

Usage:
    python3 Scripts/build_govtech_docs.py                # build all v0.2.0 papers
    python3 Scripts/build_govtech_docs.py --paper Doc/en/sovereign-by-design-v0.2.0.md
    python3 Scripts/build_govtech_docs.py --format html  # only HTML
    python3 Scripts/build_govtech_docs.py --format docx  # only DOCX

Dependencies:
    pip install markdown python-docx weasyprint
"""

import re
import sys
import argparse
from pathlib import Path
from datetime import datetime

# --- constants ---

PAPERS = [
    Path("Doc/en/sovereign-by-design-v0.2.0.md"),
    Path("Doc/de/sovereign-by-design-v0.2.0.de.md"),
]

CSS = """
/* Accessible, print-ready stylesheet for GovTech Strategy Papers */
:root {
    --primary: #1a3a5c;
    --accent: #0066cc;
    --bg: #ffffff;
    --text: #1a1a1a;
    --border: #dee2e6;
    --code-bg: #f4f6f8;
    --max-width: 860px;
}

body {
    font-family: 'Source Serif 4', Georgia, 'Times New Roman', serif;
    font-size: 1.0625rem;
    line-height: 1.75;
    color: var(--text);
    background: var(--bg);
    max-width: var(--max-width);
    margin: 0 auto;
    padding: 2rem 1.5rem 4rem;
}

h1 { font-size: 2rem; color: var(--primary); margin-top: 2rem; border-bottom: 3px solid var(--primary); padding-bottom: 0.5rem; }
h2 { font-size: 1.45rem; color: var(--primary); margin-top: 2.5rem; border-bottom: 1px solid var(--border); padding-bottom: 0.25rem; }
h3 { font-size: 1.2rem; color: var(--primary); margin-top: 1.75rem; }
h4 { font-size: 1.05rem; color: var(--primary); margin-top: 1.5rem; }

a { color: var(--accent); text-decoration: underline; }
a:hover, a:focus { color: #004499; outline: 2px solid var(--accent); outline-offset: 2px; }

table { border-collapse: collapse; width: 100%; margin: 1.5rem 0; font-size: 0.95rem; }
th { background: var(--primary); color: #fff; padding: 0.6rem 0.8rem; text-align: left; }
td { padding: 0.55rem 0.8rem; border-bottom: 1px solid var(--border); }
tr:nth-child(even) { background: #f8f9fa; }

pre, code { font-family: 'JetBrains Mono', 'Cascadia Code', 'Courier New', monospace; font-size: 0.875rem; }
pre { background: var(--code-bg); border: 1px solid var(--border); border-radius: 4px;
      padding: 1rem; overflow-x: auto; white-space: pre; line-height: 1.5; }
code { background: var(--code-bg); padding: 0.15em 0.35em; border-radius: 3px; }

blockquote { border-left: 4px solid var(--accent); margin: 1.5rem 0; padding: 0.75rem 1.25rem;
             background: #eef4ff; color: #234; border-radius: 0 4px 4px 0; }

hr { border: none; border-top: 1px solid var(--border); margin: 2.5rem 0; }

ul, ol { padding-left: 1.75rem; }
li { margin-bottom: 0.35rem; }

.metadata { background: #f8f9fa; border: 1px solid var(--border); border-radius: 6px;
            padding: 1rem 1.25rem; margin-bottom: 2rem; font-size: 0.9rem; }
.metadata p { margin: 0.2rem 0; }

@media print {
    body { max-width: 100%; padding: 1rem; font-size: 0.95rem; }
    a { color: var(--text); text-decoration: none; }
    a[href]::after { content: " (" attr(href) ")"; font-size: 0.8em; color: #555; }
    pre { page-break-inside: avoid; }
    h1, h2, h3 { page-break-after: avoid; }
}

/* Focus accessibility */
*:focus-visible { outline: 3px solid var(--accent); outline-offset: 2px; }

/* Skip link */
.skip-link { position: absolute; top: -100px; left: 0; background: var(--accent); color: #fff;
             padding: 0.5rem 1rem; z-index: 1000; text-decoration: none; border-radius: 0 0 4px 0; }
.skip-link:focus { top: 0; }
"""


def load_markdown(path: Path) -> tuple[dict, str]:
    """Load Markdown file; separate YAML front-matter from body."""
    text = path.read_text(encoding="utf-8")
    meta = {}
    body = text

    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            frontmatter = text[3:end]
            body = text[end + 4:].lstrip()
            for line in frontmatter.splitlines():
                if ":" in line:
                    key, _, val = line.partition(":")
                    meta[key.strip()] = val.strip().strip('"')

    return meta, body


def markdown_to_html(body: str, meta: dict, css: str) -> str:
    """Convert Markdown body to a complete, accessible HTML5 document."""
    try:
        import markdown
        content_html = markdown.markdown(
            body,
            extensions=["tables", "fenced_code", "toc", "attr_list"],
        )
    except ImportError:
        # Fallback: minimal regex-based conversion
        content_html = _minimal_md_to_html(body)

    title = meta.get("title", "GovTech Strategy Paper")
    author = meta.get("author", "Sebastian Graf")
    version = meta.get("version", "")
    date = meta.get("date", "")
    lang = meta.get("language", "en")

    return f"""<!DOCTYPE html>
<html lang="{lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="author" content="{author}">
    <meta name="description" content="{title} — Version {version}">
    <title>{title} v{version}</title>
    <style>{css}</style>
</head>
<body>
    <a href="#main-content" class="skip-link">Skip to main content</a>
    <div class="metadata" role="banner">
        <p><strong>{title}</strong></p>
        <p>Author: {author} &mdash; Version {version} &mdash; {date}</p>
        <p>Licence: CC-BY-4.0 | Correspondence: sebk4c@tuta.com</p>
    </div>
    <main id="main-content">
{content_html}
    </main>
    <footer role="contentinfo">
        <hr>
        <p style="font-size:0.875rem;color:#555;">
            Generated {datetime.now().strftime('%Y-%m-%d')} by build_govtech_docs.py.
            Source: <a href="https://github.com/SEBK4C/Strategy-for-City-GovTech">github.com/SEBK4C/Strategy-for-City-GovTech</a>.
            Licence: <a href="https://creativecommons.org/licenses/by/4.0/">CC-BY-4.0</a>.
        </p>
    </footer>
</body>
</html>
"""


def _minimal_md_to_html(md: str) -> str:
    """Minimal Markdown-to-HTML (fallback when python-markdown not installed)."""
    html = md
    # Headings
    html = re.sub(r"^##### (.+)$", r"<h5>\1</h5>", html, flags=re.MULTILINE)
    html = re.sub(r"^#### (.+)$", r"<h4>\1</h4>", html, flags=re.MULTILINE)
    html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)
    # Bold / italic
    html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"\*(.+?)\*", r"<em>\1</em>", html)
    # Inline code
    html = re.sub(r"`(.+?)`", r"<code>\1</code>", html)
    # Links
    html = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2">\1</a>', html)
    # HR
    html = re.sub(r"^---$", "<hr>", html, flags=re.MULTILINE)
    # Paragraphs (wrap non-tagged lines)
    lines = html.split("\n")
    result = []
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith("<"):
            result.append(f"<p>{stripped}</p>")
        else:
            result.append(line)
    return "\n".join(result)


def build_html(paper_path: Path, output_path: Path = None) -> Path:
    """Build HTML from a Markdown paper."""
    meta, body = load_markdown(paper_path)
    html = markdown_to_html(body, meta, CSS)

    if output_path is None:
        output_path = paper_path.with_suffix(".html")

    output_path.write_text(html, encoding="utf-8")
    print(f"[HTML] Written: {output_path}")
    return output_path


def build_docx(paper_path: Path, output_path: Path = None) -> Path:
    """Build DOCX from a Markdown paper using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[DOCX] python-docx not installed. Run: pip install python-docx")
        print("[DOCX] Skipping DOCX generation.")
        return None

    if output_path is None:
        output_path = paper_path.with_suffix(".docx")

    meta, body = load_markdown(paper_path)
    doc = Document()

    # Set document properties
    core_props = doc.core_properties
    core_props.author = meta.get("author", "Sebastian Graf")
    core_props.title = meta.get("title", "GovTech Strategy")
    core_props.version = meta.get("version", "0.2.0")
    core_props.language = meta.get("language", "en")

    # Style defaults
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Process Markdown lines
    for line in body.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph()
            continue
        if stripped.startswith("# ") and not stripped.startswith("## "):
            p = doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", stripped):
            text = re.sub(r"^\d+\. ", "", stripped)
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("|"): 
            # Simple table row — add as paragraph (full table parsing is complex)
            doc.add_paragraph(stripped, style="Normal")
        elif stripped == "---":
            doc.add_paragraph("―" * 60, style="Normal")
        else:
            # Strip bold/italic markers for docx body text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", stripped)
            clean = re.sub(r"\*(.+?)\*", r"\1", clean)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            doc.add_paragraph(clean)

    doc.save(str(output_path))
    print(f"[DOCX] Written: {output_path}")
    return output_path


def build_pdf(html_path: Path, output_path: Path = None) -> Path:
    """Build PDF from HTML using weasyprint."""
    try:
        from weasyprint import HTML as WeasyprintHTML
    except ImportError:
        print("[PDF] weasyprint not installed. Run: pip install weasyprint")
        print("[PDF] Skipping PDF generation.")
        return None

    if output_path is None:
        output_path = html_path.with_suffix(".pdf")

    WeasyprintHTML(filename=str(html_path)).write_pdf(str(output_path))
    print(f"[PDF] Written: {output_path}")
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Build GovTech strategy paper outputs.")
    parser.add_argument(
        "--paper",
        type=Path,
        default=None,
        help="Path to a specific Markdown paper to build.",
    )
    parser.add_argument(
        "--format",
        choices=["html", "docx", "pdf", "all"],
        default="all",
        help="Output format (default: all).",
    )
    args = parser.parse_args()

    papers = [args.paper] if args.paper else PAPERS
    papers = [p for p in papers if p.exists()]

    if not papers:
        print("No papers found. Check --paper path or run from repo root.")
        sys.exit(1)

    for paper in papers:
        print(f"\nBuilding: {paper}")
        html_path = None

        if args.format in ("html", "all"):
            html_path = build_html(paper)

        if args.format in ("docx", "all"):
            build_docx(paper)

        if args.format in ("pdf", "all") and html_path:
            build_pdf(html_path)

    print("\nBuild complete.")


if __name__ == "__main__":
    main()
