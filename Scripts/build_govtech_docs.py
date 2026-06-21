#!/usr/bin/env python3
"""Build GovTech strategy documents from Markdown into DOCX, PDF, and HTML.

The Markdown files under ``Doc/<lang>/`` are the source of truth. This script renders
each into the three published formats and writes them to ``Doc/build/<lang>/``.

Rendering strategy (in order of preference):
  1. ``pandoc`` if available — best fidelity for all of DOCX/PDF/HTML.
  2. LibreOffice CLI (``soffice``/``libreoffice``) for PDF conversion from generated DOCX.
  3. Pure-Python fallback — ``markdown`` for HTML, ``python-docx`` for DOCX. PDF is
     produced from HTML via ``weasyprint`` if native libraries are available; otherwise
     skipped with a notice.

Usage:
    uv run python Scripts/build_govtech_docs.py                 # build every paper, every format
    uv run python Scripts/build_govtech_docs.py Doc/en/foo.md   # build a single file
    uv run python Scripts/build_govtech_docs.py --formats html  # restrict output formats

Exit code is non-zero if any requested build fails.
"""
from __future__ import annotations

import argparse
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = REPO_ROOT / "Doc"
BUILD_DIR = DOC_DIR / "build"
ALL_FORMATS = ("docx", "pdf", "html")


def find_markdown_sources() -> list[Path]:
    """Return every language-versioned Markdown source under Doc/<lang>/."""
    sources: list[Path] = []
    for lang_dir in sorted(DOC_DIR.iterdir()):
        if not lang_dir.is_dir() or lang_dir.name == "build":
            continue
        sources.extend(sorted(lang_dir.glob("*.md")))
    return sources


def strip_front_matter(text: str) -> tuple[dict, str]:
    """Split YAML front-matter (between leading ``---`` fences) from the body."""
    meta: dict[str, str] = {}
    if text.startswith("---"):
        end = text.find("\n---", 3)
        if end != -1:
            raw = text[3:end].strip()
            body = text[end + 4 :].lstrip("\n")
            for line in raw.splitlines():
                m = re.match(r"^([A-Za-z0-9_-]+):\s*(.*)$", line)
                if m:
                    meta[m.group(1)] = m.group(2).strip().strip('"')
            return meta, body
    return meta, text


def have(tool: str) -> bool:
    return shutil.which(tool) is not None


def libreoffice_cmd() -> str | None:
    for tool in ("soffice", "libreoffice"):
        found = shutil.which(tool)
        if found:
            return found
    return None


def out_path(src: Path, fmt: str) -> Path:
    rel = src.relative_to(DOC_DIR)
    target = BUILD_DIR / rel.parent / f"{src.stem}.{fmt}"
    target.parent.mkdir(parents=True, exist_ok=True)
    return target


def build_with_pandoc(src: Path, fmt: str) -> bool:
    target = out_path(src, fmt)
    cmd = ["pandoc", str(src), "-o", str(target), "--standalone"]
    if fmt == "pdf":
        # Prefer weasyprint as the PDF engine (no LaTeX dependency).
        if have("weasyprint"):
            cmd += ["--pdf-engine=weasyprint"]
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
        print(f"  [pandoc] {target.relative_to(REPO_ROOT)}")
        return True
    except subprocess.CalledProcessError as exc:
        print(f"  [pandoc] FAILED {fmt} for {src.name}: {exc.stderr.strip()[:200]}")
        return False


def build_html_py(src: Path) -> bool:
    try:
        import markdown  # type: ignore
    except ImportError:
        print("  [py] skip HTML: `markdown` not installed")
        return False
    meta, body = strip_front_matter(src.read_text(encoding="utf-8"))
    html_body = markdown.markdown(body, extensions=["tables", "fenced_code", "toc"])
    title = meta.get("title", src.stem)
    doc = (
        "<!DOCTYPE html>\n<html lang=\"%s\">\n<head>\n<meta charset=\"utf-8\">\n"
        "<title>%s</title>\n<style>body{max-width:50rem;margin:2rem auto;"
        "font-family:Georgia,serif;line-height:1.5;padding:0 1rem}"
        "table{border-collapse:collapse}td,th{border:1px solid #ccc;padding:.3rem .6rem}"
        "code,pre{font-family:ui-monospace,monospace}pre{background:#f6f8fa;padding:1rem;"
        "overflow:auto}</style>\n</head>\n<body>\n%s\n</body>\n</html>\n"
    ) % (meta.get("language", "en"), title, html_body)
    target = out_path(src, "html")
    target.write_text(doc, encoding="utf-8")
    print(f"  [py] {target.relative_to(REPO_ROOT)}")
    return True


def build_docx_py(src: Path) -> bool:
    try:
        import docx  # type: ignore
        from docx.shared import Pt
    except ImportError:
        print("  [py] skip DOCX: `python-docx` not installed")
        return False
    meta, body = strip_front_matter(src.read_text(encoding="utf-8"))
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    for line in body.splitlines():
        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            document.add_heading(line.lstrip("#").strip(), level=min(level, 9))
        elif line.strip().startswith(("- ", "* ")):
            document.add_paragraph(line.strip()[2:], style="List Bullet")
        elif line.strip():
            document.add_paragraph(line)
    target = out_path(src, "docx")
    document.save(str(target))
    print(f"  [py] {target.relative_to(REPO_ROOT)}")
    return True


def build_pdf_libreoffice(src: Path) -> bool:
    tool = libreoffice_cmd()
    if not tool:
        return False

    docx_target = out_path(src, "docx")
    if not docx_target.exists() and not build_docx_py(src):
        return False

    target = out_path(src, "pdf")
    with tempfile.TemporaryDirectory(prefix="city-gov-lo-") as profile:
        cmd = [
            tool,
            "--headless",
            f"-env:UserInstallation={Path(profile).as_uri()}",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target.parent),
            str(docx_target),
        ]
        try:
            subprocess.run(cmd, check=True, capture_output=True, text=True)
            produced = docx_target.with_suffix(".pdf")
            if produced != target and produced.exists():
                produced.replace(target)
            print(f"  [libreoffice] {target.relative_to(REPO_ROOT)}")
            return True
        except subprocess.CalledProcessError as exc:
            msg = (exc.stderr or exc.stdout or "").strip()
            print(f"  [libreoffice] PDF failed for {src.name}: {msg[:200]}")
            return False


def build_pdf_py(src: Path) -> bool:
    if build_pdf_libreoffice(src):
        return True

    try:
        import weasyprint  # type: ignore
    except (ImportError, OSError) as exc:
        print(f"  [py] skip PDF: install pandoc, LibreOffice CLI, or WeasyPrint native libraries ({exc})")
        return False
    # Ensure an HTML exists, then convert.
    html_target = out_path(src, "html")
    if not html_target.exists() and not build_html_py(src):
        return False
    try:
        weasyprint.HTML(filename=str(html_target)).write_pdf(str(out_path(src, "pdf")))
        print(f"  [py] {out_path(src, 'pdf').relative_to(REPO_ROOT)}")
        return True
    except Exception as exc:  # pragma: no cover - environment dependent
        print(f"  [py] PDF failed for {src.name}: {exc}")
        return False


def build_one(src: Path, formats: tuple[str, ...]) -> bool:
    print(f"Building {src.relative_to(REPO_ROOT)}")
    ok = True
    use_pandoc = have("pandoc")
    for fmt in formats:
        if use_pandoc and build_with_pandoc(src, fmt):
            continue
        if fmt == "html":
            ok &= build_html_py(src)
        elif fmt == "docx":
            ok &= build_docx_py(src)
        elif fmt == "pdf":
            ok &= build_pdf_py(src)
    return ok


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("sources", nargs="*", help="Specific Markdown files (default: all).")
    parser.add_argument(
        "--formats", default=",".join(ALL_FORMATS),
        help="Comma-separated subset of: docx,pdf,html",
    )
    args = parser.parse_args()

    formats = tuple(f.strip() for f in args.formats.split(",") if f.strip() in ALL_FORMATS)
    sources = [Path(s).resolve() for s in args.sources] if args.sources else find_markdown_sources()
    if not sources:
        print("No Markdown sources found under Doc/<lang>/.")
        return 1

    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    if not have("pandoc"):
        print("note: pandoc not found — using fallback renderers where possible.\n")

    all_ok = True
    for src in sources:
        all_ok &= build_one(src, formats)
    print("\nDone." if all_ok else "\nDone with some skipped/failed formats (see above).")
    return 0 if all_ok else 2


if __name__ == "__main__":
    raise SystemExit(main())
