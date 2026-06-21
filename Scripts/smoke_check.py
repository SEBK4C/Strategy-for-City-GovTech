#!/usr/bin/env python3
"""Smoke-check strategy papers and generated artifacts.

Usage:
    uv run python Scripts/smoke_check.py
    uv run python Scripts/smoke_check.py --require-formats docx,html,pdf

The default gate is intentionally small enough for every self-improvement loop:
Markdown sources must look complete, EN/DE pairs must be in the same rough word-count
range, and generated DOCX/HTML artifacts must exist and broadly match Markdown length.
Use the PDF format requirement in CI or release checks where PDF tooling is installed.
"""
from __future__ import annotations

import argparse
import html.parser
import re
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = REPO_ROOT / "Doc"
BUILD_DIR = DOC_DIR / "build"
WORD_RE = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9]+(?:[-'][A-Za-zÀ-ÖØ-öø-ÿ0-9]+)?")
FRONT_MATTER_RE = re.compile(r"\A---\s*\n.*?\n---\s*\n", re.DOTALL)
URL_RE = re.compile(r"https?://\S+")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]+\)")
IMAGE_RE = re.compile(r"!\[[^\]]*\]\([^)]+\)")
CODE_FENCE_RE = re.compile(r"```.*?```", re.DOTALL)
REF_LINK_RE = re.compile(r"^\[[^\]]+\]:\s+\S+.*$", re.MULTILINE)
CONFLICT_MARKERS = ("<<<<<<<", "=======", ">>>>>>>")


class TextExtractor(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []
        self.skip_depth = 0

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"script", "style"}:
            self.skip_depth += 1

    def handle_endtag(self, tag: str) -> None:
        if tag in {"script", "style"} and self.skip_depth:
            self.skip_depth -= 1

    def handle_data(self, data: str) -> None:
        if not self.skip_depth:
            self.parts.append(data)

    def text(self) -> str:
        return " ".join(self.parts)


def strip_front_matter(text: str) -> str:
    return FRONT_MATTER_RE.sub("", text, count=1)


def markdown_text(text: str) -> str:
    text = strip_front_matter(text)
    text = CODE_FENCE_RE.sub(" ", text)
    text = IMAGE_RE.sub(" ", text)
    text = LINK_RE.sub(r"\1", text)
    text = REF_LINK_RE.sub(" ", text)
    text = URL_RE.sub(" ", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"[*_#>|~\[\](){},.;:!?/\\]", " ", text)
    return text


def word_count(text: str) -> int:
    return len(WORD_RE.findall(text))


def md_word_count(path: Path) -> int:
    return word_count(markdown_text(path.read_text(encoding="utf-8")))


def html_word_count(path: Path) -> int:
    parser = TextExtractor()
    parser.feed(path.read_text(encoding="utf-8"))
    return word_count(parser.text())


def docx_word_count(path: Path) -> int:
    import docx  # type: ignore

    document = docx.Document(str(path))
    parts: list[str] = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return word_count(" ".join(parts))


def pdf_word_count(path: Path) -> int:
    from pypdf import PdfReader  # type: ignore

    reader = PdfReader(str(path))
    text = " ".join(page.extract_text() or "" for page in reader.pages)
    return word_count(text)


def generated_path(src: Path, fmt: str) -> Path:
    rel = src.relative_to(DOC_DIR)
    return BUILD_DIR / rel.parent / f"{src.stem}.{fmt}"


def markdown_sources() -> list[Path]:
    sources: list[Path] = []
    for lang_dir in sorted(DOC_DIR.iterdir()):
        if lang_dir.is_dir() and lang_dir.name != "build":
            sources.extend(sorted(lang_dir.glob("*.md")))
    return sources


def translated_from(path: Path) -> Path | None:
    text = path.read_text(encoding="utf-8")
    if not text.startswith("---"):
        return None
    end = text.find("\n---", 3)
    if end == -1:
        return None
    for line in text[3:end].splitlines():
        key, _, value = line.partition(":")
        if key.strip() == "translated-from" and value.strip():
            candidate = (path.parent / value.strip().strip('"')).resolve()
            return candidate if candidate.exists() else None
    return None


def english_pair_for(target: Path) -> Path | None:
    explicit = translated_from(target)
    if explicit:
        return explicit
    if target.parent.name == "de" and target.name.endswith(".de.md"):
        candidate = DOC_DIR / "en" / f"{target.name.removesuffix('.de.md')}.md"
        return candidate if candidate.exists() else None
    return None


def ratio_ok(actual: int, expected: int, low: float, high: float) -> bool:
    if expected <= 0:
        return False
    ratio = actual / expected
    return low <= ratio <= high


def count_generated(path: Path, fmt: str) -> int:
    if fmt == "docx":
        return docx_word_count(path)
    if fmt == "html":
        return html_word_count(path)
    if fmt == "pdf":
        return pdf_word_count(path)
    raise ValueError(f"unsupported format: {fmt}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--require-formats",
        default="docx,html",
        help="Comma-separated generated formats that must exist: docx,html,pdf",
    )
    parser.add_argument("--min-source-words", type=int, default=200)
    parser.add_argument("--translation-min-ratio", type=float, default=0.60)
    parser.add_argument("--translation-max-ratio", type=float, default=1.45)
    parser.add_argument("--artifact-min-ratio", type=float, default=0.55)
    parser.add_argument("--artifact-max-ratio", type=float, default=1.45)
    args = parser.parse_args()

    required = {fmt.strip() for fmt in args.require_formats.split(",") if fmt.strip()}
    allowed = {"docx", "html", "pdf"}
    if bad := sorted(required - allowed):
        print(f"Unsupported format(s): {', '.join(bad)}")
        return 2

    problems = 0
    sources = markdown_sources()
    print("Smoke and sanity checks\n" + "=" * 60)
    if not sources:
        print("FAIL no Markdown sources found under Doc/<lang>/")
        return 2

    counts: dict[Path, int] = {}
    for src in sources:
        rel = src.relative_to(REPO_ROOT)
        text = src.read_text(encoding="utf-8")
        count = md_word_count(src)
        counts[src] = count
        if count < args.min_source_words:
            problems += 1
            print(f"FAIL {rel}: only {count} words")
        elif not re.search(r"^#\s+\S", strip_front_matter(text), re.MULTILINE):
            problems += 1
            print(f"FAIL {rel}: missing H1")
        elif any(marker in text for marker in CONFLICT_MARKERS):
            problems += 1
            print(f"FAIL {rel}: merge conflict marker found")
        else:
            print(f"OK   {rel}: {count} Markdown words")

    for tgt in sorted(p for p in sources if p.parent.name != "en"):
        src = english_pair_for(tgt)
        if not src:
            print(f"note {tgt.relative_to(REPO_ROOT)}: no English source pair found")
            continue
        en_count = counts.get(src) or md_word_count(src)
        tgt_count = counts[tgt]
        ratio = tgt_count / en_count if en_count else 0
        if ratio_ok(tgt_count, en_count, args.translation_min_ratio, args.translation_max_ratio):
            print(f"OK   {tgt.relative_to(REPO_ROOT)}: EN/DE word ratio {ratio:.2f}")
        else:
            problems += 1
            print(
                f"FAIL {tgt.relative_to(REPO_ROOT)}: EN/DE word ratio {ratio:.2f} "
                f"({tgt_count} vs {en_count})"
            )

    for src in sources:
        md_count = counts[src]
        for fmt in sorted(required):
            artifact = generated_path(src, fmt)
            rel = artifact.relative_to(REPO_ROOT)
            if not artifact.exists():
                problems += 1
                print(f"FAIL {rel}: missing")
                continue
            if artifact.stat().st_size == 0:
                problems += 1
                print(f"FAIL {rel}: empty")
                continue
            try:
                generated_count = count_generated(artifact, fmt)
            except Exception as exc:
                problems += 1
                print(f"FAIL {rel}: cannot read {fmt} ({exc})")
                continue
            ratio = generated_count / md_count if md_count else 0
            if ratio_ok(generated_count, md_count, args.artifact_min_ratio, args.artifact_max_ratio):
                print(f"OK   {rel}: {fmt} word ratio {ratio:.2f}")
            else:
                problems += 1
                print(
                    f"FAIL {rel}: {fmt} word ratio {ratio:.2f} "
                    f"({generated_count} vs {md_count})"
                )

    print("\n" + ("PASSED." if not problems else f"FAILED: {problems} problem(s)."))
    return 0 if not problems else 2


if __name__ == "__main__":
    sys.exit(main())
